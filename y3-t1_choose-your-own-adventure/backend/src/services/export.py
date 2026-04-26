"""ExportService — render a `Story` into DOCX or TXT bytes for download."""

from __future__ import annotations

import io
import logging
from enum import Enum

from docx import Document

from src.models.errors import StoryNotFound, UnsupportedExportFormat
from src.models.graph import ActionNode, GamebookGraph, NarrativeNode
from src.models.stories import Story
from src.repositories import StoryRepository

logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    DOCX = "docx"
    TXT = "txt"


class ExportService:
    """Walks a story's graph and emits a flat document. Pure rendering — no LLM, no storage."""

    DOCX_MIME = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    TXT_MIME = "text/plain"

    def __init__(self, story_repository: StoryRepository) -> None:
        self._stories = story_repository

    async def render(
        self,
        story_id: str,
        user_email: str,
        export_format: ExportFormat,
    ) -> tuple[bytes, str, str]:
        """Return ``(payload, mime_type, filename)`` for the rendered story."""
        story = await self._stories.get_by_id_for_user(story_id, user_email)
        if story is None:
            raise StoryNotFound(story_id)

        if export_format is ExportFormat.DOCX:
            payload = self._to_docx(story)
            return payload, self.DOCX_MIME, f"{self._safe_filename(story.name)}.docx"
        if export_format is ExportFormat.TXT:
            payload = self._to_txt(story)
            return payload, self.TXT_MIME, f"{self._safe_filename(story.name)}.txt"
        raise UnsupportedExportFormat(str(export_format))

    # ---------- Renderers ----------

    @classmethod
    def _to_txt(cls, story: Story) -> bytes:
        lines = [story.name, "=" * len(story.name), ""]
        lines.extend(cls._walk_lines(story.graph))
        return "\n".join(lines).encode("utf-8")

    @classmethod
    def _to_docx(cls, story: Story) -> bytes:
        doc = Document()
        doc.add_heading(story.name, level=1)
        for line in cls._walk_lines(story.graph):
            if line.startswith("[End]"):
                doc.add_paragraph(line, style="Intense Quote")
            elif line.startswith(">"):
                doc.add_paragraph(line, style="Quote")
            else:
                doc.add_paragraph(line)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _walk_lines(graph: GamebookGraph) -> list[str]:
        """Depth-first walk of the graph; narrative paragraphs interleaved with action lines."""
        if not graph.nodes:
            return []
        lines: list[str] = []
        visited: set[int] = set()

        def visit(node_id: int) -> None:
            if node_id in visited or node_id not in graph.node_lookup:
                return
            visited.add(node_id)
            node = graph.node_lookup[node_id]
            if isinstance(node, NarrativeNode):
                lines.append(node.data)
                if node.is_ending:
                    lines.append("[End]")
            elif isinstance(node, ActionNode):
                lines.append(f"> {node.data}")
            for child_id in node.children_ids:
                visit(child_id)

        # Start from node 0 (the canonical root of every generated story).
        roots = [n.node_id for n in graph.nodes if n.node_id == 0] or [
            graph.nodes[0].node_id
        ]
        for root in roots:
            visit(root)
        return lines

    @staticmethod
    def _safe_filename(name: str) -> str:
        cleaned = "".join(c if c.isalnum() or c in " -_" else "_" for c in name).strip()
        return cleaned or "story"
