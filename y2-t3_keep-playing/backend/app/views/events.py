import io

from django.http import StreamingHttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from ..email import notify_favourites_of_new_event
from ..models import Event, User
from ..permissions import IsOrganiser
from ..serializers import EventSerializer, UserSerializer


class EventView(APIView):
    permission_classes = [IsAuthenticated, IsOrganiser]

    def get(self, request, format=None):
        events = Event.objects.filter(
            organiser_user=request.user,
        ).order_by('date')
        serializer = EventSerializer(events, many=True)
        return Response(serializer.data)

    def post(self, request):
        data = request.data.copy()
        data['organiser_user_id'] = request.user.pk
        serializer = EventSerializer(data=data)
        if serializer.is_valid():
            event = serializer.save()
            notify_favourites_of_new_event(request.user, event)
            return Response(EventSerializer(event).data, status=status.HTTP_201_CREATED)
        return Response(
            {"error": True, "error_msg": serializer.errors},
            status=status.HTTP_400_BAD_REQUEST,
        )

    def delete(self, request, pk, format=None):
        event = get_object_or_404(Event, pk=pk)
        if event.organiser_user != request.user:
            return Response(
                {"error": True, "error_msg": "Not your event"},
                status=status.HTTP_403_FORBIDDEN,
            )
        event.delete()
        return Response({'message': 'Deleted'})

    def patch(self, request, pk, format=None):
        event = get_object_or_404(Event, pk=pk)
        if event.organiser_user != request.user:
            return Response(
                {"error": True, "error_msg": "Not your event"},
                status=status.HTTP_403_FORBIDDEN,
            )
        serializer = EventSerializer(event, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_202_ACCEPTED)
        return Response(
            {"error": True, "error_msg": serializer.error_messages},
            status=status.HTTP_400_BAD_REQUEST,
        )


class EventGetOrganiserView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk, format=None):
        event = get_object_or_404(Event, pk=pk)
        user = event.organiser_user
        serializer = UserSerializer(user, many=False)
        return Response(serializer.data)


class ExportDocx(APIView):
    permission_classes = [IsAuthenticated, IsOrganiser]

    def build_document(self, event):
        document = Document()
        document.add_heading(f"INVOICE FOR {event.name}, on {event.date}\n")
        document.add_paragraph("This is a normal style paragraph")
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.italic = True
        run.add_text("text will have italic style")
        run.add_break()
        return document

    def get(self, request, event_pk, *args, **kwargs):
        event = get_object_or_404(Event, pk=event_pk)
        document = self.build_document(event)
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = StreamingHttpResponse(
            streaming_content=buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = 'attachment;filename=Invoice.docx'
        response['Content-Encoding'] = 'UTF-8'
        return response
